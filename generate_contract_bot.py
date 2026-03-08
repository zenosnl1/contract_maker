import os
from docx import Document
import http.server
import socketserver
from core.security import access_guard, get_user_role
from core.constants import FIELDS, QUESTIONS, FlowState
from core.constants import CONTRACT_TEMPLATE, ACT_TEMPLATE, CHECKOUT_ACT_TEMPLATE, EXPENSE_CATEGORIES
from core.checkout_act import build_checkout_act
from reports.excel import build_stats_excel
from reports.finance import build_finance_report
from reports.expenses import build_expenses_report
from db.client import (
    fetch_all_contracts,
    fetch_active_contracts,
    save_contract_to_db,
    get_contract_by_code,
    insert_violation,
    fetch_contract_violations,
    calculate_close_preview,
    delete_violation,
    close_contract_full,
    insert_booking,
    fetch_active_bookings,
    insert_fixed_expense,
    insert_expense,
    fetch_fixed_expenses,
    fetch_expenses_last_30_days,
    fetch_fixed_expense_by_id,
    delete_fixed_expense,
    update_fixed_expense,
    fetch_expenses_by_month,
    fetch_contract_violations_for_period,
    fetch_penalties_by_contract_codes,
    fetch_all_expenses,
)
from telegram.ext import ApplicationBuilder
from telegram import Update
from telegram import ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ConversationHandler,
    ContextTypes,
    filters,
)
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import CallbackQueryHandler
from datetime import date, timedelta, datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


TOKEN = os.environ["BOT_TOKEN"]

# ===== Word replacement =====

async def date_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    # --- booking mode ---
    if context.user_data.get("mode") == "booking":
        return await booking_date_callback(update, context)

    # --- contract/import mode ---
    query = update.callback_query
    await query.answer()

    iso = query.data.split(":")[1]
    d = datetime.fromisoformat(iso)

    step = context.user_data["step"]
    if step >= len(FIELDS):
        return FlowState.FILLING

    field = FIELDS[step]

    context.user_data[field] = d.strftime("%d.%m.%Y")

    step += 1
    context.user_data["step"] = step

    if field == "START_DATE":
        next_day = d + timedelta(days=1)

        await query.edit_message_text(
            "📅 Выберите дату выезда:",
            reply_markup=date_keyboard(start_from=next_day),
        )
        return FlowState.FILLING

    next_field = FIELDS[step]

    if next_field == "CHECKOUT_TIME":
        await query.edit_message_text(
            "⏰ Выберите время выезда:",
            reply_markup=checkout_keyboard(),
        )
        return FlowState.FILLING

    await query.edit_message_text(QUESTIONS[next_field])
    return FlowState.FILLING



def payment_method_keyboard():
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("💵 Наличные", callback_data="PAY_CASH"),
            InlineKeyboardButton("🏦 Банковский перевод", callback_data="PAY_BANK"),
        ]
    ])


def invoice_keyboard():
    return InlineKeyboardMarkup([
        [
            InlineKeyboardButton("✅ Да", callback_data="INVOICE_YES"),
            InlineKeyboardButton("❌ Нет", callback_data="INVOICE_NO"),
        ]
    ])

def checkout_keyboard():
    buttons = [
        [
            InlineKeyboardButton("09:00", callback_data="CHECKOUT:09:00"),
            InlineKeyboardButton("12:00", callback_data="CHECKOUT:12:00"),
        ],
        [
            InlineKeyboardButton("15:00", callback_data="CHECKOUT:15:00"),
            InlineKeyboardButton("18:00", callback_data="CHECKOUT:18:00"),
        ],
    ]

    return InlineKeyboardMarkup(buttons)

def start_keyboard(user):

    role = get_user_role(user)

    buttons = []

    # --- admin only ---
    if role == "admin":
        buttons += [
            [InlineKeyboardButton("▶️ Начать оформление", callback_data="START_FLOW")],
            [InlineKeyboardButton("📥 Импорт договора", callback_data="MENU_IMPORT")],
            [InlineKeyboardButton("✏️ Управление договором", callback_data="MENU_EDIT")],
            [InlineKeyboardButton("🚨 Нарушения", callback_data="MENU_VIOLATIONS_MENU")],
            [InlineKeyboardButton("📌 Брони", callback_data="MENU_BOOKINGS")],
            [InlineKeyboardButton("💸 Расходы", callback_data="MENU_EXPENSES")],
        ]

    # --- admin + viewer ---
    if role in ("admin", "viewer"):
        buttons += [
            [InlineKeyboardButton("📊 Отчёты", callback_data="MENU_STATS_MENU")],
            [InlineKeyboardButton("👥 Текущие жильцы", callback_data="MENU_ACTIVE")],
        ]

    return InlineKeyboardMarkup(buttons)

def parse_price(value: str) -> float | None:
    """
    Принимает:
      12
      12.5
      12,5
    Возвращает float или None.
    """

    txt = value.strip().replace(",", ".")

    try:
        amount = float(txt)
    except ValueError:
        return None

    if amount <= 0:
        return None

    return round(amount, 2)

async def back_to_fixed_menu(update, context):

    query = update.callback_query
    await query.answer()

    return await show_fixed_expenses_menu(update, context)

async def stats_expenses_callback(update, context):

    query = update.callback_query
    await query.answer()

    rows = fetch_all_expenses()

    await query.edit_message_text("📊 Формирую отчёт по расходам...")

    path = build_expenses_report(rows)

    with open(path, "rb") as f:
        await query.message.reply_document(f)

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU


async def fixed_expenses_menu_callback(update, context):

    query = update.callback_query
    await query.answer()

    return await show_fixed_expenses_menu(update, context)

async def fixed_expense_list(update, context):

    query = update.callback_query
    await query.answer()

    rows = fetch_fixed_expenses()

    if not rows:
        await query.edit_message_text(
            "📭 Регулярных расходов пока нет.",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_EXPENSES")],
            ])
        )
        return FlowState.FIXED_EXPENSE_MENU

    lines = ["📅 Регулярные расходы:\n"]

    total_all = 0.0

    for r in rows:

        total_all += float(r["total_price"])

        lines.append(
            f"{r['id']}) {r['item_name']}\n"
            f"   🔢 Кол-во: {r['quantity']}\n"
            f"   💶 Цена: {float(r['unit_price']):.2f} €\n"
            f"   💸 Итого: {float(r['total_price']):.2f} €\n"
        )

    lines.append("━━━━━━━━━━━━")
    lines.append(f"💸 Общий регулярный расход: {total_all:.2f} €")

    await query.edit_message_text(
        "\n".join(lines),
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_EXPENSES")],
        ])
    )

    return FlowState.FIXED_EXPENSE_MENU



async def show_fixed_expenses_menu(update, context):

    msg = update.message or update.callback_query.message

    await msg.reply_text(
        "📅 Регулярные расходы\n\nВыберите действие:",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("📋 Просмотреть", callback_data="FIXED_LIST")],
            [InlineKeyboardButton("➕ Создать", callback_data="FIXED_CREATE")],
            [InlineKeyboardButton("✏️ Отредактировать", callback_data="FIXED_EDIT")],
            [InlineKeyboardButton("🗑 Удалить", callback_data="FIXED_DELETE")],
            [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_EXPENSES")],
        ])
    )

    return FlowState.FIXED_EXPENSE_MENU

async def fixed_expense_delete_start(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text("Введите ID регулярного расхода для удаления:")

    return FlowState.FIXED_EXPENSE_DELETE_SELECT

async def fixed_expense_delete_enter(update, context):

    txt = update.message.text.strip()

    if not txt.isdigit():
        await update.message.reply_text("Введите числовой ID.")
        return FlowState.FIXED_EXPENSE_DELETE_SELECT

    fid = int(txt)

    row = fetch_fixed_expense_by_id(fid)

    if not row:
        await update.message.reply_text("❌ Расход с таким ID не найден.")
        return FlowState.FIXED_EXPENSE_DELETE_SELECT

    delete_fixed_expense(fid)

    await update.message.reply_text("🗑 Регулярный расход удалён.")

    return await show_fixed_expenses_menu(update, context)

async def fixed_expense_edit_select(update, context):

    txt = update.message.text.strip()

    if not txt.isdigit():
        await update.message.reply_text("Введите числовой ID.")
        return FlowState.FIXED_EXPENSE_EDIT_SELECT

    fid = int(txt)

    row = fetch_fixed_expense_by_id(fid)

    if not row:
        await update.message.reply_text("❌ Расход не найден.")
        return FlowState.FIXED_EXPENSE_EDIT_SELECT

    context.user_data["fixed_mode"] = "edit"
    context.user_data["fixed_expense"] = dict(row)  # содержит id

    await update.message.reply_text(
        f"Редактирование:\n\n"
        f"{row['id']}) {row['item_name']}\n"
        f"Кол-во: {row['quantity']}\n"
        f"Цена: {row['unit_price']} €\n\n"
        "Введите новое количество:"
    )

    return FlowState.FIXED_EXPENSE_CREATE_QTY

async def fixed_expense_edit_qty(update, context):

    txt = update.message.text.strip()

    if not txt.isdigit():
        await update.message.reply_text("Введите число.")
        return FlowState.FIXED_EXPENSE_CREATE_QTY

    context.user_data["fixed_expense"]["quantity"] = int(txt)

    await update.message.reply_text("Введите новую цену:")

    return FlowState.FIXED_EXPENSE_CREATE_PRICE

async def fixed_expense_edit_price(update, context):

    txt = update.message.text.strip()

    price = parse_price(txt)

    if price is None:
        await update.message.reply_text("Введите цену, например 12,5")
        return FlowState.FIXED_EXPENSE_CREATE_PRICE

    fe = context.user_data["fixed_expense"]

    total = round(fe["quantity"] * price, 3)

    payload = {
        "quantity": fe["quantity"],
        "unit_price": price,
        "total_price": total,
    }

    update_fixed_expense(fe["id"], payload)

    await update.message.reply_text("✅ Регулярный расход обновлён.")

    return await show_fixed_expenses_menu(update, context)


async def fixed_expense_create_start(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data["fixed_mode"] = "create"
    context.user_data["fixed_expense"] = {}

    await query.edit_message_text("Введите название предмета:")

    return FlowState.FIXED_EXPENSE_CREATE_NAME

async def fixed_expense_name_enter(update, context):

    context.user_data["fixed_expense"]["item_name"] = update.message.text.strip()

    await update.message.reply_text("Введите количество:")

    return FlowState.FIXED_EXPENSE_CREATE_QTY

async def fixed_expense_qty_enter(update, context):

    txt = update.message.text.strip()

    if not txt.isdigit():
        await update.message.reply_text("Введите количество цифрами.")
        return FlowState.FIXED_EXPENSE_CREATE_QTY

    context.user_data["fixed_expense"]["quantity"] = int(txt)

    await update.message.reply_text("Введите цену за единицу:")

    return FlowState.FIXED_EXPENSE_CREATE_PRICE

async def fixed_expense_price_enter(update, context):

    txt = update.message.text.strip()
    
    price = parse_price(txt)
    
    if price is None:
        await update.message.reply_text(
            "Введите цену, например: 12,5 или 12.50"
        )
        return FlowState.FIXED_EXPENSE_CREATE_PRICE
    
    unit_price = price

    fe = context.user_data["fixed_expense"]
    fe["unit_price"] = unit_price

    total = round(fe["quantity"] * unit_price, 3)
    fe["total_price"] = total

    payload = {
        "item_name": fe["item_name"],
        "quantity": fe["quantity"],
        "unit_price": fe["unit_price"],
        "total_price": total,
    }
    
    mode = context.user_data.get("fixed_mode")
    
    if mode == "edit":
        update_fixed_expense(fe["id"], payload)
        await update.message.reply_text("✏️ Регулярный расход обновлён.")
    else:
        insert_fixed_expense(payload)
        await update.message.reply_text("✅ Регулярный расход сохранён.")

    await update.message.reply_text(
        "✅ Регулярный расход сохранён:\n\n"
        f"📦 {fe['item_name']}\n"
        f"🔢 Кол-во: {fe['quantity']}\n"
        f"💶 Цена: {fe['unit_price']} €\n"
        f"💸 Итого: {total:.3f} €"
    )

    context.user_data.pop("fixed_expense", None)
    context.user_data.pop("fixed_mode", None)

    return await show_fixed_expenses_menu(update, context)


async def fixed_expense_edit_start(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text("Введите ID регулярного расхода для редактирования:")

    return FlowState.FIXED_EXPENSE_EDIT_SELECT


async def back_to_expenses_menu(update, context):

    query = update.callback_query
    await query.answer()

    return await expenses_menu_callback(update, context)


async def expenses_menu_callback(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text(
        "💸 Расходы\n\nВыберите действие:",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("➕ Учесть расход", callback_data="EXPENSE_ADD")],
            [InlineKeyboardButton("📆 Расходы за последние 30 дней", callback_data="EXPENSE_LAST30")],
            [InlineKeyboardButton("🧾 Расходы за месяц", callback_data="EXPENSE_MONTH")],
            [InlineKeyboardButton("⏳ Регулярные расходы", callback_data="EXPENSE_FIXED")],
            [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_MENU")],
        ])
    )

    return FlowState.EXPENSES_MENU

async def expenses_month_pick(update, context):

    query = update.callback_query
    await query.answer()

    today = date.today()

    buttons = []

    for i in range(12):
        y = today.year
        m = today.month - i

        while m <= 0:
            m += 12
            y -= 1

        label = f"{m:02}.{y}"
        key = f"{y}-{m:02}"

        buttons.append([
            InlineKeyboardButton(
                label,
                callback_data=f"EXPENSE_MONTH_SHOW:{key}",
            )
        ])

    buttons.append([
        InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_EXPENSES"),
    ])
    
    await query.edit_message_text(
        "📅 Выберите месяц:",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

    return FlowState.EXPENSE_MONTH_PICK

async def expenses_month_show(update, context):

    query = update.callback_query
    await query.answer()

    ym = query.data.split(":")[1]  # YYYY-MM

    year, month = ym.split("-")

    rows = fetch_expenses_by_month(year, month)

    if not rows:
        await query.edit_message_text(
            f"📭 За {month}.{year} расходов нет.",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("⬅️ Назад", callback_data="EXPENSE_MONTH")],
            ]),
        )
        return FlowState.EXPENSE_MONTH_PICK

    rows = sorted(rows, key=lambda r: r["expense_date"])

    lines = [f"📅 Расходы за {month}.{year}\n"]

    total = 0.0
    total_cash = 0.0
    total_company = 0.0

    for r in rows:

        amount = float(r["amount"])
        total += amount

        if r["payment_method"] == "cash":
            total_cash += amount
        else:
            total_company += amount

        dt = datetime.fromisoformat(r["expense_date"]).strftime("%d.%m.%Y")

        pay = "Наличные" if r["payment_method"] == "cash" else "С фирмы"

        desc = r.get("description") or r.get("comment") or "—"

        lines.append(
            f"📅 {dt}\n"
            f"🛒 {desc}\n"
            f"💶 {amount:.2f} €\n"
            f"💳 {pay}\n"
        )

    lines.append("━━━━━━━━━━━━")
    lines.append(f"💸 Итого: {total:.2f} €")
    lines.append(f"🏢 С фирмы: {total_company:.2f} €")
    lines.append(f"💵 Наличными: {total_cash:.2f} €")

    await query.edit_message_text(
        "\n".join(lines),
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("⬅️ Назад", callback_data="EXPENSE_MONTH")],
        ]),
    )

    return FlowState.EXPENSE_MONTH_PICK


async def fixed_expense_delete_enter(update, context):

    txt = update.message.text.strip()

    if not txt.isdigit():
        await update.message.reply_text("Введите числовой ID.")
        return FlowState.FIXED_EXPENSE_DELETE_SELECT

    fid = int(txt)

    row = fetch_fixed_expense_by_id(fid)

    if not row:
        await update.message.reply_text("❌ Расход с таким ID не найден.")
        return FlowState.FIXED_EXPENSE_DELETE_SELECT

    context.user_data["delete_fixed_expense"] = row

    await update.message.reply_text(
        f"Удалить этот регулярный расход?\n\n"
        f"{row['id']}) {row['item_name']}\n"
        f"Кол-во: {row['quantity']}\n"
        f"Цена: {row['unit_price']} €\n"
        f"Итого: {row['total_price']} €",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("✅ Удалить", callback_data="FIXED_DELETE_CONFIRM"),
                InlineKeyboardButton("❌ Отмена", callback_data="BACK_TO_FIXED"),
            ]
        ])
    )

    return FlowState.FIXED_EXPENSE_DELETE_CONFIRM

async def fixed_expense_delete_confirm(update, context):

    query = update.callback_query
    await query.answer()

    row = context.user_data.pop("delete_fixed_expense", None)

    if not row:
        return await show_fixed_expenses_menu(update, context)

    delete_fixed_expense(row["id"])

    await query.edit_message_text("🗑 Регулярный расход удалён.")

    return await show_fixed_expenses_menu(update, context)

async def back_to_fixed(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data.pop("delete_fixed_expense", None)

    return await show_fixed_expenses_menu(update, context)



async def expenses_last30_list(update, context):

    query = update.callback_query
    await query.answer()

    rows = fetch_expenses_last_30_days()

    rows = sorted(
        rows,
        key=lambda r: r["expense_date"],
    )

    if not rows:
        await query.edit_message_text(
            "📭 За последние 30 дней расходов не найдено.",
            reply_markup=InlineKeyboardMarkup([
                [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_EXPENSES")],
            ]),
        )
        return FlowState.EXPENSES_MENU

    lines = ["📆 Расходы за последние 30 дней\n"]

    total = 0.0
    total_cash = 0.0
    total_company = 0.0

    for r in rows:

        amount = float(r["amount"])

        total += amount
        
        if r["payment_method"] == "cash":
            total_cash += amount
        else:
            total_company += amount
    
        raw_date = r["expense_date"]
    
        try:
            dt_obj = datetime.fromisoformat(raw_date)
            dt = dt_obj.strftime("%d.%m.%Y")
        except Exception:
            dt = raw_date
    
        pay = "Наличные" if r["payment_method"] == "cash" else "С фирмы"
    
        desc = r.get("description") or r.get("comment") or "—"
    
        lines.append(
            f"📅 {dt}\n"
            f"🛒 {desc}\n"
            f"💶 {float(r['amount']):.2f} €\n"
            f"💳 {pay}\n"
        )


    lines.append("━━━━━━━━━━━━")
    lines.append(f"💸 Итого за 30 дней: {total:.2f} €")
    lines.append(f"🏢 С фирмы: {total_company:.2f} €")
    lines.append(f"💵 Наличными: {total_cash:.2f} €")

    await query.edit_message_text(
        "\n".join(lines),
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_EXPENSES")],
        ]),
    )

    return FlowState.EXPENSES_MENU

async def expense_payment_chosen(update, context):

    query = update.callback_query
    await query.answer()

    if query.data == "EXP_PAY_COMPANY":
        method = "company"
    else:
        method = "cash"

    exp = context.user_data.get("expense")

    if not exp:
        await query.edit_message_text("⚠️ Ошибка: нет данных расхода.")
        return FlowState.MENU

    payload = {
        "expense_date": exp["date"],
        "amount": round(float(exp["amount"]), 2),
        "payment_method": method,
        "description": exp.get("description"),
        "comment": None,
    }

    insert_expense(payload)

    await query.edit_message_text("✅ Расход сохранён.")

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    context.user_data.pop("expense", None)

    return FlowState.MENU


async def expense_add_start(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data["expense"] = {}

    await query.edit_message_text("Введите сумму расхода:")

    return FlowState.EXPENSE_ENTER_AMOUNT


async def expense_enter_amount(update, context):

    txt = update.message.text.strip()

    amount = parse_price(txt)

    if amount is None:
        await update.message.reply_text(
            "Введите сумму. Пример:\n"
            "12\n12.5\n12,5"
        )
        return FlowState.EXPENSE_ENTER_AMOUNT

    context.user_data["expense"]["amount"] = amount

    await update.message.reply_text(
        "Выберите дату:",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("📅 Сегодня", callback_data="EXP_DATE_TODAY"),
                InlineKeyboardButton("✍️ Вписать дату вручную", callback_data="EXP_DATE_MANUAL"),
            ]
        ])
    )

    return FlowState.EXPENSE_DATE_CHOICE



async def expense_date_today(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data["expense"]["date"] = date.today().isoformat()

    msg = update.message or update.callback_query.message

    await msg.reply_text(
        "🧾 Что было куплено? (можно написать несколько позиций)"
    )
    
    return FlowState.EXPENSE_DESCRIPTION

async def expense_date_manual_start(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text("Введите дату (ДД.ММ.ГГГГ):")

    return FlowState.EXPENSE_DATE_MANUAL

async def expense_date_manual_enter(update, context):

    txt = update.message.text.strip()

    try:
        d = datetime.strptime(txt, "%d.%m.%Y").date()
    except ValueError:
        await update.message.reply_text("Формат ДД.ММ.ГГГГ")
        return FlowState.EXPENSE_DATE_MANUAL

    context.user_data["expense"]["date"] = d.isoformat()

    msg = update.message or update.callback_query.message

    await msg.reply_text(
        "🧾 Что было куплено? (можно написать несколько позиций)"
    )
    
    return FlowState.EXPENSE_DESCRIPTION

async def expense_description_enter(update, context):

    txt = update.message.text.strip()

    if len(txt) < 3:
        await update.message.reply_text("Опишите покупку текстом.")
        return FlowState.EXPENSE_DESCRIPTION

    context.user_data["expense"]["description"] = txt

    await update.message.reply_text(
        "Как был оплачен расход?",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("🏢 С фирмы", callback_data="EXP_PAY_COMPANY"),
                InlineKeyboardButton("💵 Наличными", callback_data="EXP_PAY_CASH"),
            ]
        ])
    )

    return FlowState.EXPENSE_PAYMENT_METHOD


def date_keyboard(days=30, start_from=None):

    if start_from:
        base = start_from
    else:
        base = date.today()

    buttons = []

    for i in range(days):
        d = base + timedelta(days=i)
        buttons.append([
            InlineKeyboardButton(
                d.strftime("%d.%m.%Y"),
                callback_data=f"DATE:{d.isoformat()}",
            )
        ])

    return InlineKeyboardMarkup(buttons)

def skip_keyboard():
    return InlineKeyboardMarkup(
        [[InlineKeyboardButton("⏭ Пропустить", callback_data="SKIP")]]
    )

async def bookings_menu_callback(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text(
        "📌 Брони\n\nВыберите действие:",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("➕ Создать бронь", callback_data="BOOKING_CREATE")],
            [InlineKeyboardButton("📋 Текущие брони", callback_data="BOOKING_LIST")],
            [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_MENU")],
        ])
    )

    return FlowState.BOOKING_MENU

async def booking_create_start(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data["mode"] = "booking"
    context.user_data["booking"] = {}

    await query.edit_message_text("Введите номер помещения:")

    return FlowState.BOOKING_CREATE_FLAT

async def booking_flat_enter(update, context):

    context.user_data["booking"]["flat_number"] = update.message.text.strip()
    await update.message.reply_text("Введите имя клиента:")

    return FlowState.BOOKING_CREATE_NAME


async def booking_name_enter(update, context):

    context.user_data["booking"]["client_name"] = update.message.text.strip()
    await update.message.reply_text("Введите телефон:")

    return FlowState.BOOKING_CREATE_PHONE


async def booking_phone_enter(update, context):

    context.user_data["booking"]["client_number"] = update.message.text.strip()

    await update.message.reply_text("Введите цену за ночь (€):")

    return FlowState.BOOKING_CREATE_PRICE

async def booking_price_enter(update, context):

    txt = update.message.text.strip()

    if not txt.isdigit():
        await update.message.reply_text("Введите сумму цифрами.")
        return FlowState.BOOKING_CREATE_PRICE

    context.user_data["booking"]["price_per_day"] = int(txt)

    await update.message.reply_text(
        "📅 Выберите дату заезда:",
        reply_markup=date_keyboard(),
    )

    return FlowState.BOOKING_CREATE_START


async def booking_date_callback(update, context):

    query = update.callback_query
    await query.answer()

    iso = query.data.split(":")[1]
    d = datetime.fromisoformat(iso).date()

    booking = context.user_data.setdefault("booking", {})

    # ---------- заезд ----------
    if "start_date" not in booking:

        booking["start_date"] = d.isoformat()

        await query.edit_message_text(
            "📅 Выберите дату выезда:",
            reply_markup=booking_end_keyboard(d + timedelta(days=1)),
        )

        return FlowState.BOOKING_CREATE_END

    # ---------- выезд ----------
    booking["end_date"] = d.isoformat()

    return await booking_finish(update, context)


def booking_end_keyboard(start_from):

    base_kb = date_keyboard(start_from=start_from)

    buttons = list(base_kb.inline_keyboard)

    buttons.append([
        InlineKeyboardButton("❓ Пока неизвестно", callback_data="BOOKING_END_UNKNOWN"),
    ])

    return InlineKeyboardMarkup(buttons)

async def booking_end_unknown(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data["booking"]["end_date"] = None

    return await booking_finish(update, context)

async def booking_finish(update, context):

    b = context.user_data["booking"]

    # --- даты ---
    start = datetime.fromisoformat(b["start_date"]).date()

    if b["end_date"]:
        end = datetime.fromisoformat(b["end_date"]).date()
        nights = (end - start).days
    else:
        end = None
        nights = None

    price = b["price_per_day"]

    total = nights * price if nights is not None else None

    # --- payload в БД ---
    payload = {
        "flat_number": b["flat_number"],
        "client_name": b["client_name"],
        "client_number": b["client_number"],
        "start_date": start.isoformat(),
        "price_per_day": price,
        "status": "active",
    }
    
    # optional поля
    if end:
        payload["end_date"] = end.isoformat()
        payload["nights"] = nights
        payload["total_price"] = total
    
    insert_booking(payload)

    # --- красиво пользователю ---
    start_txt = start.strftime("%d.%m.%Y")

    end_txt = end.strftime("%d.%m.%Y") if end else "❓"

    text = (
        "📌 Бронь создана:\n\n"
        f"🏠 Помещение: {b['flat_number']}\n"
        f"👤 Клиент: {b['client_name']}\n"
        f"📞 Телефон: {b['client_number']}\n"
        f"📅 Заезд: {start_txt}\n"
        f"📅 Выезд: {end_txt}\n"
        f"🌙 Ночей: {nights if nights is not None else '—'}\n"
        f"💶 Цена/ночь: {price} €\n"
        f"💰 Сумма: {total if total is not None else '—'} €"
    )

    msg = update.message or update.callback_query.message

    await msg.reply_text(text)

    await msg.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU


async def booking_list_callback(update, context):

    query = update.callback_query
    await query.answer()

    rows = fetch_active_bookings()

    if not rows:
        await query.edit_message_text("📭 Активных броней нет.")
        await query.message.reply_text(
            "Главное меню:",
            reply_markup=start_keyboard(update.effective_user),
        )
        return FlowState.MENU

    lines = ["📌 Текущие брони:\n"]

    for r in rows:

        start_txt = datetime.fromisoformat(r["start_date"]).strftime("%d.%m.%Y")

        end_txt = (
            datetime.fromisoformat(r["end_date"]).strftime("%d.%m.%Y")
            if r["end_date"]
            else "❓"
        )

        nights = r.get("nights") or "—"
        total = r.get("total_price") or "—"

        sep = "━━━━━━━━━━━━━━━━━━━━"

        lines.append(
            f"\n{sep}\n\n"
            f"🏠 {r['flat_number']}\n"
            f"👤 {r['client_name']}\n"
            f"📞 {r['client_number']}\n"
            f"📅 {start_txt} → {end_txt}\n\n"
            f"🌙 Ночей: {nights}\n"
            f"💶 Цена/ночь: {r['price_per_day']} €\n"
            f"💰 Сумма: {total} €\n"
            f"\n{sep}\n"
        )

    await query.edit_message_text("\n".join(lines))

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU


def replace_everywhere(doc, data):
    for p in doc.paragraphs:
        process_paragraph(p, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p, data)


def process_paragraph(p, data):
    text = p.text
    keys_used = [k for k in data if f"{{{{{k}}}}}" in text]

    if not keys_used:
        return

    for r in p.runs:
        r.text = ""

    i = 0
    while i < len(text):
        replaced = False
        for k in keys_used:
            ph = f"{{{{{k}}}}}"
            if text.startswith(ph, i):
                run = p.add_run(data[k])
                run.bold = True
                i += len(ph)
                replaced = True
                break

        if not replaced:
            run = p.add_run(text[i])
            i += 1


def add_page_numbers(doc):

    section = doc.sections[0]
    footer = section.footer

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# ======================================================
# Violations flow
# ======================================================

VIOLATION_REASONS = {
    "smoking": "Курение в помещении",
    "noise": "Нарушение режима тишины",
    "damage": "Повреждение помещения или оснащения",
    "dirty": "Помещение оставлено в ненадлежащем состоянии",
}

# ======================================================
# Violations menu
# ======================================================

async def violations_menu_callback(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text(
        "🚨 Нарушения\n\nВыберите действие:",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("➕ Отметить нарушение", callback_data="VIOL_ADD")],
            [InlineKeyboardButton("🗑 Удалить нарушение", callback_data="VIOL_DELETE")],
            [InlineKeyboardButton("⬅️ Назад", callback_data="BACK_TO_MENU")],
        ])
    )

    return FlowState.MENU

async def close_show_violations(update, context):

    c = context.user_data["edit_contract"]

    violations = fetch_contract_violations(c["contract_code"])

    if not violations:
        context.user_data["close_total_penalty"] = 0
        return await close_show_preview(update, context)

    total = sum(int(v["amount"]) for v in violations)

    context.user_data["close_total_penalty"] = total

    lines = [
        "📋 Найдены нарушения:\n"
    ]

    for v in violations:
        label = VIOLATION_REASONS.get(v["violation_type"], v["violation_type"])
        lines.append(f"• {label} — {v['amount']}€")

    lines.append(f"\n💶 Итого удержание: {total} €")
    lines.append("\nЗакрыть договор?")

    keyboard = InlineKeyboardMarkup([
        [
            InlineKeyboardButton("✅ Подтвердить", callback_data="CLOSE_CONFIRM"),
            InlineKeyboardButton("❌ Отмена", callback_data="CLOSE_CANCEL"),
        ]
    ])

    if update.callback_query:
        await update.callback_query.edit_message_text("\n".join(lines), reply_markup=keyboard)
    else:
        await update.message.reply_text("\n".join(lines), reply_markup=keyboard)

    return FlowState.CLOSE_CONFIRM_VIOLATIONS

async def violation_cancel(update, context):

    query = update.callback_query
    await query.answer()

    # очищаем временные данные
    context.user_data.pop("violation_contract", None)
    context.user_data.pop("violation_reason", None)
    context.user_data.pop("violation_amount", None)

    await query.edit_message_text(
        "❌ Отменено.",
    )

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU

async def back_to_menu_callback(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU


async def violation_start_callback(update, context):

    query = update.callback_query
    await query.answer()

    rows = fetch_active_contracts()

    if not rows:
        await query.edit_message_text(
            "Сейчас нет активных жильцов."
        )
        return FlowState.MENU

    buttons = []

    for r in rows:
        label = f"{r['flat_number']} — {r['client_name']}"
        buttons.append([
            InlineKeyboardButton(
                label,
                callback_data=f"VIOL_FLAT:{r['contract_code']}",
            )
        ])

    await query.edit_message_text(
        "Выберите помещение:",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

    return FlowState.VIOLATION_SELECT_FLAT


async def violation_select_flat(update, context):

    query = update.callback_query
    await query.answer()

    code = query.data.split(":")[1]

    contract = get_contract_by_code(code)

    context.user_data["violation_contract"] = contract

    buttons = []

    for k, label in VIOLATION_REASONS.items():
        buttons.append([
            InlineKeyboardButton(
                label,
                callback_data=f"VIOL_REASON:{k}",
            )
        ])

    await query.edit_message_text(
        "Укажите причину нарушения:",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

    return FlowState.VIOLATION_SELECT_REASON


async def violation_select_reason(update, context):

    query = update.callback_query
    await query.answer()

    key = query.data.split(":")[1]

    context.user_data["violation_reason"] = key

    await query.edit_message_text(
        "Введите сумму (€), которая будет удержана из депозита:"
    )

    return FlowState.VIOLATION_ENTER_AMOUNT


async def violation_enter_amount(update, context):

    val = update.message.text.strip()

    if not val.isdigit():
        await update.message.reply_text("Введите сумму цифрами.")
        return FlowState.VIOLATION_ENTER_AMOUNT

    context.user_data["violation_amount"] = int(val)

    c = context.user_data["violation_contract"]

    await update.message.reply_text(
        "📋 Проверьте данные:\n\n"
        f"🏠 Помещение: {c['flat_number']}\n"
        f"👤 Клиент: {c['client_name']}\n"
        f"🚨 Причина: {VIOLATION_REASONS[context.user_data['violation_reason']]}\n"
        f"💶 Сумма удержания: {val} €\n\n"
        "Продолжить?",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("✅ Да", callback_data="VIOL_CONFIRM"),
                InlineKeyboardButton("❌ Отмена", callback_data="VIOL_CANCEL"),
            ]
        ])
    )

    return FlowState.VIOLATION_CONFIRM


async def violation_confirm(update, context):

    query = update.callback_query
    await query.answer()

    c = context.user_data["violation_contract"]

    payload = {
        "contract_code": c["contract_code"],
        "flat_number": c["flat_number"],
        "violation_type": context.user_data["violation_reason"],
        "amount": context.user_data["violation_amount"],
        "description": None,
    }

    insert_violation(payload)

    await query.edit_message_text("✅ Нарушение сохранено.")

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU

async def import_flow_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await require_admin(update):
        return FlowState.MENU
    
    query = update.callback_query
    await query.answer()

    context.user_data.clear()
    context.user_data["step"] = 0
    context.user_data["mode"] = "import"

    await query.edit_message_text(
        "📥 Импорт договора.\n\n"
        "Введите номер помещения:"
    )

    return FlowState.FILLING

# ======================================================
# Violations delete flow
# ======================================================

async def violation_delete_start(update, context):

    query = update.callback_query
    await query.answer()

    rows = fetch_active_contracts()

    if not rows:
        await query.edit_message_text("Нет активных договоров.")
        return FlowState.MENU

    buttons = []

    for r in rows:
        label = f"{r['flat_number']} — {r['client_name']}"
        buttons.append([
            InlineKeyboardButton(
                label,
                callback_data=f"VIOL_DEL_FLAT:{r['contract_code']}",
            )
        ])

    await query.edit_message_text(
        "Выберите помещение:",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

    return FlowState.VIOLATION_DELETE_SELECT_FLAT

async def violation_delete_select_flat(update, context):

    query = update.callback_query
    await query.answer()

    code = query.data.split(":")[1]

    context.user_data["violation_delete_code"] = code

    violations = fetch_contract_violations(code)

    if not violations:
        await query.edit_message_text(
            "По этому договору нет нарушений.",
            reply_markup=start_keyboard(update.effective_user),
        )
        return FlowState.MENU

    buttons = []

    for v in violations:
        label = f"{VIOLATION_REASONS[v['violation_type']]} — {v['amount']}€"
        buttons.append([
            InlineKeyboardButton(
                label,
                callback_data=f"VIOL_DEL_ITEM:{v['id']}",
            )
        ])

    await query.edit_message_text(
        "Выберите нарушение для удаления:",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

    return FlowState.VIOLATION_DELETE_SELECT_ITEM

async def violation_delete_item(update, context):

    query = update.callback_query
    await query.answer()

    vid = query.data.split(":")[1]

    delete_violation(vid)

    await query.edit_message_text("✅ Нарушение удалено.")

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU


def generate_docs(data):
    safe = data["CLIENT_NAME"].replace(" ", "_")

    outputs = []

    for tpl, prefix in [
        (CONTRACT_TEMPLATE, "contract"),
        (ACT_TEMPLATE, "act"),
    ]:
        doc = Document(tpl)
        replace_everywhere(doc, data)
        add_page_numbers(doc)

        fname = f"{prefix}_{safe}.docx"
        doc.save(fname)
        outputs.append(fname)

    return outputs

# ===== Telegram flow =====

async def start(update, context):

    user = update.effective_user

    role = get_user_role(user)

    # неизвестен → просим телефон
    if not role:

        kb = ReplyKeyboardMarkup(
            [[KeyboardButton("📱 Поделиться номером", request_contact=True)]],
            resize_keyboard=True,
            one_time_keyboard=True,
        )

        await update.message.reply_text(
            "🔐 Для доступа к боту необходимо подтвердить номер телефона.",
            reply_markup=kb,
        )

        return FlowState.WAIT_PHONE

    context.user_data.clear()

    await update.message.reply_text(
        "👋 Главное меню:",
        reply_markup=start_keyboard(user),
    )

    return FlowState.MENU

async def phone_received(update, context):

    contact = update.message.contact

    phone = normalize_phone(contact.phone_number)

    user = update.effective_user

    # подставляем временно
    user.phone_number = phone

    role = get_user_role(user)

    if not role:
        await update.message.reply_text("⛔ Доступ запрещён.")
        return ConversationHandler.END

    await update.message.reply_text(
        "✅ Доступ подтверждён.",
        reply_markup=start_keyboard(user),
    )

    return FlowState.MENU


async def stop(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    await update.message.reply_text(
        "🛑 Процесс заполнения остановлен.",
        reply_markup=start_keyboard(update.effective_user),
    )
    return FlowState.MENU

async def back(update: Update, context: ContextTypes.DEFAULT_TYPE):
    step = context.user_data.get("step", 0)

    if step <= 0:
        await update.message.reply_text(
            "Вы уже в начале. Введите значение или используйте /stop."
        )
        return FlowState.FILLING

    step -= 1
    context.user_data["step"] = step

    field = FIELDS[step]

    await update.message.reply_text(
        f"⬅️ Возврат назад.\n\n{QUESTIONS[field]}"
    )

    return FlowState.FILLING

async def status(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not context.user_data:
        await update.message.reply_text("Пока ничего не введено.")
        return FlowState.FILLING

    lines = ["📋 Текущие данные:"]

    for f in FIELDS:
        if f in context.user_data:
            lines.append(f"• {f}: {context.user_data[f]}")

    await update.message.reply_text("\n".join(lines))
    return FlowState.FILLING

async def stats_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    try:
        rows = fetch_all_contracts()
    except Exception as e:
        print("🔥 STATS ERROR:", repr(e))
        await query.edit_message_text("⚠️ Ошибка получения данных.", reply_markup=None)
        return FlowState.MENU

    if not rows:
        await query.edit_message_text("Пока нет договоров.", reply_markup=None)
        return FlowState.MENU

    # --------------------------------------
    # подтягиваем штрафы из violations
    # --------------------------------------

    codes = [
        r["contract_code"]
        for r in rows
        if r.get("contract_code")
    ]

    penalties_map = fetch_penalties_by_contract_codes(codes)

    for r in rows:
        r["penalties"] = penalties_map.get(
            r.get("contract_code"),
            0,
        )

    # --------------------------------------

    path = build_stats_excel(rows)

    await query.edit_message_text("📊 Формирую статистику…", reply_markup=None)

    with open(path, "rb") as f:
        await query.message.reply_document(f)
    
    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU

async def stats_menu_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    await query.edit_message_text(
        "Выберите тип отчёта:",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("📊 Общий отчет", callback_data="STATS_GENERAL")],
            [InlineKeyboardButton("💰 Финансовый отчёт", callback_data="STATS_FINANCE")],
            [InlineKeyboardButton("💸 Отчёт по расходам", callback_data="STATS_EXPENSES")],
        ])
    )

    return FlowState.MENU

async def stats_finance_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    rows = fetch_all_contracts()

    await query.edit_message_text("💰 Формирую финансовый отчёт...")

    path = build_finance_report(rows)

    with open(path, "rb") as f:
        await query.message.reply_document(f)

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU

async def active_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    query = update.callback_query
    await query.answer()

    try:
        rows = fetch_active_contracts()
        def flat_key(r):
            try:
                return int(r["flat_number"])
            except Exception:
                return r["flat_number"]
        
        rows = sorted(rows, key=flat_key)

    except Exception:
        await query.edit_message_text("⚠️ Ошибка получения данных.", reply_markup=None)
        return FlowState.MENU

    if not rows:
        await query.edit_message_text("Сейчас жильцов нет.", reply_markup=None)
        return FlowState.MENU

    lines = ["👥 Текущие жильцы:\n"]

    today = date.today()

    for r in rows:

        try:
            start = datetime.fromisoformat(r["start_date"]).date()
            end = datetime.fromisoformat(r["end_date"]).date()
        
            nights = int(r["nights"])
            price = int(r["price_per_day"])
            total = int(r["total_price"])
        
            lived_nights = (today - start).days
            lived_nights = max(0, min(lived_nights, nights))
        
            remaining_nights = nights - lived_nights
        
            earned = lived_nights * price
        
            preview = calculate_close_preview(
                contract_code=r["contract_code"],
                actual_checkout_date=today,
                early_checkout=True,
                initiator="tenant",
                early_reason=None,
                manual_refund=None,
            )
            
            deposit = int(r.get("deposit") or 0)

            refund_today = preview["refund"] - deposit
            refund_today = max(0, refund_today)

            extra_due = preview["extra_due"]
            penalties = preview["penalties"]

        except Exception as e:
            print("🔥 ACTIVE ROW ERROR:", r)
            print(e)
            continue
    
        separator = "━━━━━━━━━━━━━━━━━━━━"

        lines.append(
            f"\n{separator}\n\n"
        
            f"🏠 {r['flat_number']}\n"
            f"👤 {r['client_name']}\n"
            f"📞 {r['client_number']}\n"
            f"📅 {r['start_date']} → {r['end_date']}\n\n"
        
            f"✅ Прожито: {lived_nights} ночей / {earned} €\n"
            f"⏳ Осталось: {remaining_nights} ночей\n"
            f"💰 Депозит: {deposit} €\n"
            f"💸 Возврат при выезде сегодня (без депозита): {refund_today} €\n"
            f"⚠️ Удержания/долг: {extra_due + penalties} €\n"
        
            f"\n{separator}\n"
        )

    await query.edit_message_text("\n".join(lines), reply_markup=None)

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU

async def start_flow_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    context.user_data.clear()
    context.user_data["mode"] = "normal"
    context.user_data["step"] = 0

    await query.edit_message_text(
        "📄 Начинаем создание договора.\n\n"
        + QUESTIONS[FIELDS[0]]
    , reply_markup=None)

    return FlowState.FILLING

async def checkout_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    time_val = query.data.replace("CHECKOUT:", "")

    step = context.user_data["step"]
    field = FIELDS[step]  # CHECKOUT_TIME

    context.user_data[field] = time_val

    step += 1
    context.user_data["step"] = step

    next_field = FIELDS[step]

    await query.edit_message_text(QUESTIONS[next_field], reply_markup=None)
    return FlowState.FILLING

async def skip_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    step = context.user_data["step"]
    field = FIELDS[step]

    context.user_data[field] = "-----"

    step += 1
    context.user_data["step"] = step

    next_field = FIELDS[step]

    if next_field in ["CLIENT_ADDRESS", "CLIENT_MAIL"]:
        await query.edit_message_text(
            QUESTIONS[next_field],
            reply_markup=skip_keyboard(),
        )
        return FlowState.FILLING
    
    await query.edit_message_text(QUESTIONS[next_field], reply_markup=None)
    return FlowState.FILLING

async def handle_answer(update: Update, context: ContextTypes.DEFAULT_TYPE):

    mode = context.user_data.get("mode", "normal")
    step = context.user_data["step"]
    field = FIELDS[step]

    text = update.message.text.strip()

    # ---------- ВАЛИДАЦИЯ ----------

    mode = context.user_data.get("mode", "normal")
    
    if field in ["START_DATE", "END_DATE"] and mode == "import":
    
        try:
            datetime.strptime(text, "%d.%m.%Y")
        except ValueError:
            await update.message.reply_text(
                "❌ Формат даты должен быть ДД.ММ.ГГГГ"
            )
            return FlowState.FILLING

    if field in ["MAX_PEOPLE_DAY", "MAX_PEOPLE_NIGHT"]:
        if not text.isdigit():
            await update.message.reply_text(
                "❌ Введите число, например: 4"
            )
            return FlowState.FILLING

    if field == "PRICE_PER_DAY":
        if not text.isdigit():
            await update.message.reply_text(
                "❌ Введите цену цифрами, например: 25"
            )
            return FlowState.FILLING

    if field == "DEPOSIT":
        if not text.isdigit():
            await update.message.reply_text(
                "❌ Введите депозит цифрами, например: 80"
            )
            return FlowState.FILLING

    # ---------- СОХРАНЯЕМ ----------

    context.user_data[field] = text

    # ---------- АВТОРАСЧЁТ СУММЫ ----------

    if field == "PRICE_PER_DAY":

        start = datetime.strptime(context.user_data["START_DATE"], "%d.%m.%Y")
        end = datetime.strptime(context.user_data["END_DATE"], "%d.%m.%Y")

        nights = (end - start).days
        total = nights * int(text)

        context.user_data["TOTAL_PRICE"] = str(total)

        await update.message.reply_text(
            f"💶 {nights} ночей × {text} € = {total} €"
        )

    # ====== PAYMENT FLOW ======

    if field == "DEPOSIT":

        await update.message.reply_text(
            "💳 Как производится оплата?",
            reply_markup=payment_method_keyboard(),
        )
    
        return FlowState.PAYMENT_METHOD



    # ---------- ДВИГАЕМСЯ ВПЕРЁД ----------

    step += 1
    context.user_data["step"] = step

    # ---------- ЕСЛИ ЕСТЬ СЛЕДУЮЩИЙ ШАГ ----------

    if step < len(FIELDS):

        next_field = FIELDS[step]

        if next_field == "START_DATE":
            if mode == "import":
                await update.message.reply_text(
                    "Введите дату заезда (ДД.ММ.ГГГГ):"
                )
            else:
                await update.message.reply_text(
                    "📅 Выберите дату заезда:",
                    reply_markup=date_keyboard(),
                )
        
            return FlowState.FILLING

        if next_field == "END_DATE":
            if mode == "import":
                await update.message.reply_text(
                    "Введите дату выезда (ДД.ММ.ГГГГ):"
                )
            else:
                await update.message.reply_text(
                    "📅 Выберите дату выезда:",
                    reply_markup=date_keyboard(),
                )
        
            return FlowState.FILLING


        if next_field == "CHECKOUT_TIME":
            await update.message.reply_text(
                "⏰ Выберите время выезда:",
                reply_markup=checkout_keyboard(),
            )
            return FlowState.FILLING

        if next_field in ["CLIENT_ADDRESS", "CLIENT_MAIL"]:
            await update.message.reply_text(
                QUESTIONS[next_field],
                reply_markup=skip_keyboard(),
            )
            return FlowState.FILLING

        await update.message.reply_text(QUESTIONS[next_field])
        return FlowState.FILLING

    # ---------- ФИНАЛ: ГЕНЕРИРУЕМ ДОКУМЕНТЫ ----------

    files = generate_docs(context.user_data)

    context.user_data["_generated_files"] = files

    await update.message.reply_text(
        "📄 Документы готовы.\n\n"
        "Сохранить договор в базе данных?",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("💾 Да", callback_data="SAVE_DB"),
                InlineKeyboardButton("❌ Нет", callback_data="SKIP_DB"),
            ]
        ])
    )

    return FlowState.CONFIRM_SAVE

class Handler(http.server.SimpleHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.end_headers()
        self.wfile.write(b"OK")

async def payment_method_callback(update, context):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    if query.data == "PAY_CASH":

        context.user_data["PAYMENT_METHOD"] = "cash"
        context.user_data["INVOICE_ISSUED"] = False
        context.user_data["INVOICE_NUMBER"] = None

        return await continue_after_payment(update, context)

    # bank
    context.user_data["PAYMENT_METHOD"] = "bank_transfer"

    await query.edit_message_text(
        "Был ли выставлен счёт?",
        reply_markup=invoice_keyboard(),
    )

    return FlowState.PAYMENT_INVOICE

async def invoice_choice_callback(update, context):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    if query.data == "INVOICE_NO":

        context.user_data["INVOICE_ISSUED"] = False
        context.user_data["INVOICE_NUMBER"] = None

        return await continue_after_payment(update, context)

    context.user_data["INVOICE_ISSUED"] = True

    await query.edit_message_text("Введите номер счёта:")

    return FlowState.PAYMENT_INVOICE_NUMBER

async def invoice_number_enter(update, context):

    context.user_data["INVOICE_NUMBER"] = update.message.text.strip()

    return await continue_after_payment(update, context)
    
async def continue_after_payment(update, context):

    # <<< ТУТ двигаем шаг >>>
    step = context.user_data["step"] + 1
    context.user_data["step"] = step

    # если FIELDS закончились — финал
    if step >= len(FIELDS):
        files = generate_docs(context.user_data)
        context.user_data["_generated_files"] = files

        msg = update.message or update.callback_query.message

        await msg.reply_text(
            "📄 Документы готовы.\n\n"
            "Сохранить договор в базе данных?",
            reply_markup=InlineKeyboardMarkup([
                [
                    InlineKeyboardButton("💾 Да", callback_data="SAVE_DB"),
                    InlineKeyboardButton("❌ Нет", callback_data="SKIP_DB"),
                ]
            ])
        )

        return FlowState.CONFIRM_SAVE

    next_field = FIELDS[step]

    msg = update.message or update.callback_query.message
    await msg.reply_text(QUESTIONS[next_field])

    return FlowState.FILLING


async def save_db_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END

    if await require_admin(update):
        return FlowState.MENU
    
    query = update.callback_query
    await query.answer()

    rows = fetch_fixed_expenses()

    fixed_sum = round(
        sum(float(r["total_price"]) for r in rows),
        2,
    )
    
    context.user_data["FIXED_PER_BOOKING"] = fixed_sum

    save_contract_to_db(
        context.user_data,
        context.user_data["_generated_files"],
    )

    for fpath in context.user_data["_generated_files"]:
        with open(fpath, "rb") as f:
            await query.message.reply_document(f)

    await query.edit_message_text("💾 Сохранено.", reply_markup=None)
    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    context.user_data.clear()

    return FlowState.MENU

async def skip_db_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    for fpath in context.user_data["_generated_files"]:
        with open(fpath, "rb") as f:
            await query.message.reply_document(f)

    await query.edit_message_text("Не Сохранено.", reply_markup=None)
    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    context.user_data.clear()

    return FlowState.MENU

async def edit_menu_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    rows = fetch_active_contracts()

    buttons = []

    for r in rows:
        label = f"{r['flat_number']} — {r['client_name']}"
        buttons.append([
            InlineKeyboardButton(
                label,
                callback_data=f"EDIT_ACTIVE:{r['contract_code']}",
            )
        ])

    buttons.append([
        InlineKeyboardButton("✍️ Ввести номер вручную", callback_data="EDIT_MANUAL"),
    ])

    await query.edit_message_text(
        "Выберите активный договор:",
        reply_markup=InlineKeyboardMarkup(buttons),
    )

    return FlowState.EDIT_SELECT_ACTIVE

async def edit_select_active(update, context):

    if await access_guard(update):
        return ConversationHandler.END
    
    query = update.callback_query
    await query.answer()

    code = query.data.split(":")[1]

    contract = get_contract_by_code(code)

    if not contract:
        await query.edit_message_text("❌ Договор не найден.")
        return FlowState.MENU

    context.user_data["edit_contract"] = contract
    context.user_data["close_contract_code"] = code

    await query.edit_message_text(
        "Что сделать с договором?",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("📄 Информация", callback_data="SHOW_CONTRACT_INFO")],
            [InlineKeyboardButton("🏁 Завершить договор", callback_data="CLOSE_CONTRACT")]
        ])
    )

    return FlowState.EDIT_ACTION

async def edit_manual_enter(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text("Введите номер договора:")

    return FlowState.EDIT_ENTER_CODE


async def edit_enter_code_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):

    code = update.message.text.strip()

    contract = get_contract_by_code(code)

    if not contract:
        await update.message.reply_text("❌ Договор не найден. Попробуйте снова.")
        return FlowState.EDIT_ENTER_CODE

    if contract.get("is_closed"):
        await update.message.reply_text(
            "⚠️ Этот договор уже закрыт.\n\n"
            "Выберите другой номер договора.",
        )
        return FlowState.EDIT_ENTER_CODE

    context.user_data["edit_contract"] = contract
    context.user_data["close_contract_code"] = code

    await update.message.reply_text(
        "Что сделать с договором?",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("📄 Информация", callback_data="SHOW_CONTRACT_INFO")],
            [InlineKeyboardButton("🏁 Завершить договор", callback_data="CLOSE_CONTRACT")]
        ])
    )

    return FlowState.EDIT_ACTION

async def show_contract_info_callback(update, context):

    query = update.callback_query
    await query.answer()

    contract = context.user_data.get("edit_contract")

    if not contract:
        await query.edit_message_text("❌ Нет данных договора.")
        return FlowState.MENU

    text = format_contract_view(contract)

    await query.edit_message_text(text)

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU


# ======================================================
# Extended close flow with act generation
# ======================================================

async def close_select_initiator(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data["early_checkout"] = True

    await query.edit_message_text(
        "Кто инициировал досрочный выезд?",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("👤 Клиент", callback_data="EARLY_TENANT"),
                InlineKeyboardButton("🏠 Арендодатель", callback_data="EARLY_LANDLORD"),
            ]
        ])
    )

    return FlowState.CLOSE_SELECT_INITIATOR


async def close_initiator_chosen(update, context):

    query = update.callback_query
    await query.answer()

    initiator = query.data.replace("EARLY_", "").lower()

    context.user_data["early_initiator"] = initiator

    if initiator == "tenant":

        await query.edit_message_text(
            "Укажите причину досрочного выезда:"
        )

        return FlowState.CLOSE_ENTER_EARLY_REASON

    # landlord
    await query.edit_message_text(
        "Как определить возврат?",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("📊 Рассчитать автоматически", callback_data="LANDLORD_AUTO"),
                InlineKeyboardButton("✍️ Ввести вручную", callback_data="LANDLORD_MANUAL"),
            ]
        ])
    )

    return FlowState.CLOSE_LANDLORD_REFUND_MODE


async def close_enter_early_reason(update, context):

    context.user_data["early_reason"] = update.message.text.strip()

    return await ask_close_date(update, context)


async def close_landlord_refund_mode(update, context):

    query = update.callback_query
    await query.answer()

    mode = query.data

    if mode == "LANDLORD_MANUAL":

        await query.edit_message_text("Введите сумму возврата (€):")
        return FlowState.CLOSE_ENTER_MANUAL_REFUND

    context.user_data["manual_refund"] = None

    return await ask_close_date(update, context)

async def close_enter_manual_refund(update, context):

    txt = update.message.text.strip()

    if not txt.isdigit():
        await update.message.reply_text("Введите сумму цифрами.")
        return FlowState.CLOSE_ENTER_MANUAL_REFUND

    context.user_data["manual_refund"] = int(txt)

    return await ask_close_date(update, context)

async def ask_close_date(update, context):

    kb = InlineKeyboardMarkup([
        [
            InlineKeyboardButton("Сегодня", callback_data="CLOSE_TODAY"),
            InlineKeyboardButton("Ввести вручную", callback_data="CLOSE_MANUAL"),
        ]
    ])

    text = "📅 Укажите фактическую дату выезда:"

    if update.callback_query:
        await update.callback_query.edit_message_text(text, reply_markup=kb)
    else:
        await update.message.reply_text(text, reply_markup=kb)

    return FlowState.CLOSE_ENTER_DATE

async def close_show_preview(update, context):

    c = context.user_data["edit_contract"]

    result = calculate_close_preview(
        contract_code=c["contract_code"],
        actual_checkout_date=context.user_data["actual_end_date"],
        early_checkout=context.user_data.get("early_checkout", False),
        initiator=context.user_data.get("early_initiator"),
        early_reason=context.user_data.get("early_reason"),
        manual_refund=context.user_data.get("manual_refund"),
    )

    context.user_data["close_calc"] = result

    lines = [
        "📋 Предпросмотр закрытия:\n",
        f"Прожито ночей: {result['lived_nights']}",
        f"Непрожито → {result['unused']}€",
        f"Штрафы: {result['penalties']}€",
        f"Возврат: {result['refund']}€",
        f"Долг клиента: {result['extra_due']}€",
        "",
        "Закрыть договор и сформировать акт?"
    ]

    kb = InlineKeyboardMarkup([
        [
            InlineKeyboardButton("✅ Подтвердить", callback_data="CLOSE_FINAL_CONFIRM"),
            InlineKeyboardButton("❌ Отмена", callback_data="CLOSE_CANCEL"),
        ]
    ])

    if update.callback_query:
        await update.callback_query.edit_message_text("\n".join(lines), reply_markup=kb)
    else:
        await update.message.reply_text("\n".join(lines), reply_markup=kb)

    return FlowState.CLOSE_PREVIEW_ACT

async def close_contract_start(update: Update, context: ContextTypes.DEFAULT_TYPE):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text(
        "Досрочное завершение?",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("Да", callback_data="CLOSE_EARLY_YES"),
                InlineKeyboardButton("Нет", callback_data="CLOSE_EARLY_NO"),
            ]
        ])
    )

    return FlowState.CLOSE_IS_EARLY

async def close_cancel(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text(
        "❌ Закрытие договора отменено.",
    )

    await query.message.reply_text(
        "Главное меню:",
        reply_markup=start_keyboard(update.effective_user),
    )

    return FlowState.MENU

async def close_early_yes(update, context):

    query = update.callback_query
    await query.answer()

    return await close_select_initiator(update, context)

async def close_early_no(update, context):

    query = update.callback_query
    await query.answer()

    contract = context.user_data["edit_contract"]

    planned_end = datetime.fromisoformat(contract["end_date"]).date()

    context.user_data["actual_end_date"] = planned_end
    context.user_data["early_checkout"] = False
    context.user_data.pop("early_initiator", None)
    context.user_data.pop("early_reason", None)
    context.user_data.pop("manual_refund", None)

    return await close_show_violations(update, context)

async def close_today(update, context):

    query = update.callback_query
    await query.answer()

    context.user_data["actual_end_date"] = datetime.today().date()

    return await close_show_preview(update, context)

async def close_manual(update, context):

    query = update.callback_query
    await query.answer()

    await query.edit_message_text("Введите дату закрытия (ДД.ММ.ГГГГ):")

    return FlowState.CLOSE_ENTER_DATE

async def close_receive_date(update, context):

    try:
        d = datetime.strptime(update.message.text, "%d.%m.%Y").date()
    except ValueError:
        await update.message.reply_text("❌ Формат: ДД.ММ.ГГГГ")
        return FlowState.CLOSE_ENTER_DATE

    context.user_data["actual_end_date"] = d

    return await close_show_preview(update, context)

async def require_admin(update):

    role = get_user_role(update.effective_user)

    if role != "admin":
        msg = update.message or update.callback_query.message
        await msg.reply_text("⛔ У вас нет прав для этого действия.")
        return True

    return False


async def finalize_close(update, context):

    c = context.user_data["edit_contract"]
    
    if c.get("is_closed"):
        await update.message.reply_text("⚠️ Договор уже закрыт.")
        return FlowState.MENU

    result = close_contract_full(
        contract_code=c["contract_code"],
        actual_checkout_date=context.user_data["actual_end_date"],
        early_checkout=context.user_data.get("early_checkout"),
        initiator=context.user_data.get("early_initiator"),
        early_reason=context.user_data.get("early_reason"),
        manual_refund=context.user_data.get("manual_refund"),
    )

    # 🔥 заново читаем договор из БД
    contract = get_contract_by_code(c["contract_code"])

    violations = fetch_contract_violations_for_period(
        contract_code=contract["contract_code"],
        start_date=contract["start_date"],
        actual_end_date=contract["actual_checkout_date"],
    )

    safe_code = contract["contract_code"].replace("/", "_")

    path = build_checkout_act(
        template_path=CHECKOUT_ACT_TEMPLATE,
        output_path=f"checkout_act_{safe_code}.docx",
        contract=contract,
        violations=violations,
    )

    msg = update.message or update.callback_query.message

    with open(path, "rb") as f:
        await msg.reply_document(f)

    await msg.reply_text(
        "✅ Договор закрыт и акт сформирован.",
        reply_markup=start_keyboard(update.effective_user),
    )

    context.user_data.clear()

    return FlowState.MENU
    
def format_contract_view(c: dict) -> str:

    def v(x):
        return x if x not in [None, "", "-----"] else "-"

    lines = [
        "📄 Договор\n",

        f"🆔 Код: {v(c.get('contract_code'))}",
        f"🏠 Помещение: {v(c.get('flat_number'))}",

        "",

        f"👤 Клиент: {v(c.get('client_name'))}",
        f"📄 Документ: {v(c.get('client_id'))}",
        f"📞 Телефон: {v(c.get('client_number'))}",
        f"📧 Email: {v(c.get('client_mail'))}",
        f"🏠 Адрес: {v(c.get('client_address'))}",

        "",

        f"📅 Заезд: {v(c.get('start_date'))}",
        f"📅 Плановый выезд: {v(c.get('end_date'))}",
        f"📅 Фактический выезд: {v(c.get('actual_checkout_date'))}",
        f"⏰ Время выезда: {v(c.get('checkout_time'))}",

        "",

        f"🌙 Ночей: {v(c.get('nights'))}",
        f"💶 Цена/ночь: {v(c.get('price_per_day'))} €",
        f"💰 Общая сумма: {v(c.get('total_price'))} €",
        f"💳 Депозит: {v(c.get('deposit'))} €",

        "",
    ]

    # ----- Статус договора -----
    status = "Завершён" if c.get("is_closed") else "Активен"
    lines.append(f"📦 Статус договора: {status}")

    # ----- Payment -----
    pm_raw = c.get("payment_method")

    if pm_raw == "cash":
        pm = "Наличные"
    elif pm_raw == "bank_transfer":
        pm = "Банковский перевод"
    else:
        pm = "—"

    lines.append(f"💳 Способ оплаты: {pm}")

    if c.get("invoice_issued"):
        lines.append(f"📄 Счёт: {v(c.get('invoice_number'))}")

    return "\n".join(x for x in lines if x)


# ===== main =====

WEBHOOK_PATH = "/webhook"
PORT = int(os.environ.get("PORT", 10000))
PUBLIC_URL = os.environ.get("PUBLIC_URL")  # будем задать в Render

def main():
    port = int(os.environ.get("PORT", 10000))
    public_url = os.environ.get("PUBLIC_URL")

    if not public_url:
        raise RuntimeError("PUBLIC_URL env var is not set")

    webhook_url = public_url.rstrip("/") + WEBHOOK_PATH

    print("🌍 Webhook URL:", webhook_url)

    app = ApplicationBuilder().token(TOKEN).build()
    
    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            FlowState.MENU: [
                CallbackQueryHandler(start_flow_callback, pattern="^START_FLOW$"),
                CallbackQueryHandler(import_flow_callback, pattern="^MENU_IMPORT$"),
                CallbackQueryHandler(bookings_menu_callback, pattern="^MENU_BOOKINGS$"),
                CallbackQueryHandler(booking_list_callback, pattern="^BOOKING_LIST$"),
                CallbackQueryHandler(expenses_menu_callback, pattern="^MENU_EXPENSES$"),
                CallbackQueryHandler(stats_expenses_callback, pattern="^STATS_EXPENSES$"),
                CallbackQueryHandler(violations_menu_callback, pattern="^MENU_VIOLATIONS_MENU$"),
                CallbackQueryHandler(violation_start_callback, pattern="^VIOL_ADD$"),
                CallbackQueryHandler(violation_delete_start, pattern="^VIOL_DELETE$"),
                CallbackQueryHandler(back_to_menu_callback, pattern="^BACK_TO_MENU$"),
                CallbackQueryHandler(edit_menu_callback, pattern="^MENU_EDIT$"),
                CallbackQueryHandler(stats_menu_callback, pattern="^MENU_STATS_MENU$"),
                CallbackQueryHandler(stats_callback, pattern="^STATS_GENERAL$"),
                CallbackQueryHandler(stats_finance_callback, pattern="^STATS_FINANCE$"),
                CallbackQueryHandler(active_callback, pattern="^MENU_ACTIVE$"),
            ],
    
            FlowState.FILLING: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, handle_answer),
    
                CallbackQueryHandler(date_callback, pattern="^DATE:"),
                CallbackQueryHandler(checkout_callback, pattern="^CHECKOUT:"),
                CallbackQueryHandler(skip_callback, pattern="^SKIP$"),
    
                CommandHandler("back", back),
                CommandHandler("status", status),
                CommandHandler("stop", stop),
            ],
            FlowState.PAYMENT_METHOD: [
                CallbackQueryHandler(payment_method_callback, pattern="^PAY_"),
            ],
            
            FlowState.PAYMENT_INVOICE: [
                CallbackQueryHandler(invoice_choice_callback, pattern="^INVOICE_"),
            ],
            
            FlowState.PAYMENT_INVOICE_NUMBER: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, invoice_number_enter),
            ],
            FlowState.CONFIRM_SAVE: [
                CallbackQueryHandler(save_db_callback, pattern="^SAVE_DB$"),
                CallbackQueryHandler(skip_db_callback, pattern="^SKIP_DB$"),
            ],
            FlowState.EDIT_SELECT_ACTIVE: [
                CallbackQueryHandler(edit_select_active, pattern="^EDIT_ACTIVE:"),
                CallbackQueryHandler(edit_manual_enter, pattern="^EDIT_MANUAL$"),
            ],
            FlowState.EDIT_ENTER_CODE: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, edit_enter_code_handler),
            ],
        
            FlowState.EDIT_ACTION: [
                CallbackQueryHandler(close_contract_start, pattern="^CLOSE_CONTRACT$"),
                CallbackQueryHandler(show_contract_info_callback, pattern="^SHOW_CONTRACT_INFO$"),
            ],
            FlowState.CLOSE_IS_EARLY: [
                CallbackQueryHandler(close_early_yes, pattern="^CLOSE_EARLY_YES$"),
                CallbackQueryHandler(close_early_no, pattern="^CLOSE_EARLY_NO$"),
            ],
        
            FlowState.CLOSE_ENTER_DATE: [
                CallbackQueryHandler(close_today, pattern="^CLOSE_TODAY$"),
                CallbackQueryHandler(close_manual, pattern="^CLOSE_MANUAL$"),
                MessageHandler(filters.TEXT & ~filters.COMMAND, close_receive_date),
            ],
            FlowState.VIOLATION_SELECT_FLAT: [
                CallbackQueryHandler(violation_select_flat, pattern="^VIOL_FLAT:"),
            ],
            
            FlowState.VIOLATION_SELECT_REASON: [
                CallbackQueryHandler(violation_select_reason, pattern="^VIOL_REASON:"),
            ],
            
            FlowState.VIOLATION_ENTER_AMOUNT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, violation_enter_amount),
            ],
            
            FlowState.VIOLATION_CONFIRM: [
                CallbackQueryHandler(violation_confirm, pattern="^VIOL_CONFIRM$"),
                CallbackQueryHandler(violation_cancel, pattern="^VIOL_CANCEL$"),
            ],
            FlowState.VIOLATION_DELETE_SELECT_FLAT: [
                CallbackQueryHandler(violation_delete_select_flat, pattern="^VIOL_DEL_FLAT:"),
            ],
            
            FlowState.VIOLATION_DELETE_SELECT_ITEM: [
                CallbackQueryHandler(violation_delete_item, pattern="^VIOL_DEL_ITEM:"),
            ],
            FlowState.CLOSE_CONFIRM_VIOLATIONS: [
                CallbackQueryHandler(close_show_preview, pattern="^CLOSE_CONFIRM$"),
                CallbackQueryHandler(close_cancel, pattern="^CLOSE_CANCEL$"),
            ],
            FlowState.CLOSE_SELECT_INITIATOR: [
                CallbackQueryHandler(close_initiator_chosen, pattern="^EARLY_"),
            ],
            
            FlowState.CLOSE_ENTER_EARLY_REASON: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, close_enter_early_reason),
            ],
            
            FlowState.CLOSE_LANDLORD_REFUND_MODE: [
                CallbackQueryHandler(close_landlord_refund_mode, pattern="^LANDLORD_"),
            ],
            
            FlowState.CLOSE_ENTER_MANUAL_REFUND: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, close_enter_manual_refund),
            ],
            
            FlowState.CLOSE_PREVIEW_ACT: [
                CallbackQueryHandler(finalize_close, pattern="^CLOSE_FINAL_CONFIRM$"),
                CallbackQueryHandler(close_cancel, pattern="^CLOSE_CANCEL$"),
            ],
            FlowState.WAIT_PHONE: [
                MessageHandler(filters.CONTACT, phone_received),
            ],
            FlowState.BOOKING_MENU: [
                CallbackQueryHandler(booking_create_start, pattern="^BOOKING_CREATE$"),
                CallbackQueryHandler(booking_list_callback, pattern="^BOOKING_LIST$"),
                CallbackQueryHandler(back_to_menu_callback, pattern="^BACK_TO_MENU$"),
            ],
            
            FlowState.BOOKING_CREATE_FLAT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, booking_flat_enter),
            ],
            
            FlowState.BOOKING_CREATE_NAME: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, booking_name_enter),
            ],
            
            FlowState.BOOKING_CREATE_PHONE: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, booking_phone_enter),
            ],
            FlowState.BOOKING_CREATE_PRICE: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, booking_price_enter),
            ],
            
            FlowState.BOOKING_CREATE_START: [
                CallbackQueryHandler(date_callback, pattern="^DATE:"),
            ],
            
            FlowState.BOOKING_CREATE_END: [
                CallbackQueryHandler(date_callback, pattern="^DATE:"),
                CallbackQueryHandler(booking_end_unknown, pattern="^BOOKING_END_UNKNOWN$"),
            ],
            FlowState.EXPENSES_MENU: [
                CallbackQueryHandler(expense_add_start, pattern="^EXPENSE_ADD$"),
                CallbackQueryHandler(back_to_expenses_menu, pattern="^BACK_TO_EXPENSES$"),
                CallbackQueryHandler(expenses_last30_list, pattern="^EXPENSE_LAST30$"),
                CallbackQueryHandler(expenses_month_pick, pattern="^EXPENSE_MONTH$"),
                CallbackQueryHandler(fixed_expenses_menu_callback, pattern="^EXPENSE_FIXED$"),
                CallbackQueryHandler(back_to_menu_callback, pattern="^BACK_TO_MENU$"),
            ],
            FlowState.EXPENSE_ENTER_AMOUNT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, expense_enter_amount),
            ],
            FlowState.EXPENSE_DATE_CHOICE: [
                CallbackQueryHandler(expense_date_today, pattern="^EXP_DATE_TODAY$"),
                CallbackQueryHandler(expense_date_manual_start, pattern="^EXP_DATE_MANUAL$"),
            ],
            
            FlowState.EXPENSE_DATE_MANUAL: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, expense_date_manual_enter),
            ],
            
            FlowState.EXPENSE_DESCRIPTION: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, expense_description_enter),
            ],
            FlowState.EXPENSE_PAYMENT_METHOD: [
                CallbackQueryHandler(expense_payment_chosen, pattern="^EXP_PAY_"),
            ],
            FlowState.FIXED_EXPENSE_MENU: [
                CallbackQueryHandler(fixed_expenses_menu_callback, pattern="^EXPENSE_FIXED$"),
                CallbackQueryHandler(back_to_expenses_menu, pattern="^BACK_TO_EXPENSES$"),
                CallbackQueryHandler(fixed_expense_list, pattern="^FIXED_LIST$"),
                CallbackQueryHandler(fixed_expense_create_start, pattern="^FIXED_CREATE$"),
                CallbackQueryHandler(fixed_expense_edit_start, pattern="^FIXED_EDIT$"),
                CallbackQueryHandler(fixed_expense_delete_start, pattern="^FIXED_DELETE$"),
            ],
            FlowState.FIXED_EXPENSE_CREATE_NAME: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, fixed_expense_name_enter),
            ],
            
            FlowState.FIXED_EXPENSE_CREATE_QTY: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, fixed_expense_qty_enter),
            ],
            
            FlowState.FIXED_EXPENSE_CREATE_PRICE: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, fixed_expense_price_enter),
            ],
            FlowState.FIXED_EXPENSE_DELETE_SELECT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, fixed_expense_delete_enter),
            ],
            
            FlowState.FIXED_EXPENSE_EDIT_SELECT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, fixed_expense_edit_select),
            ],
            FlowState.FIXED_EXPENSE_DELETE_CONFIRM: [
                CallbackQueryHandler(fixed_expense_delete_confirm, pattern="^FIXED_DELETE_CONFIRM$"),
                CallbackQueryHandler(back_to_fixed, pattern="^BACK_TO_FIXED$"),
            ],
            FlowState.EXPENSE_MONTH_PICK: [
                CallbackQueryHandler(back_to_expenses_menu, pattern="^BACK_TO_EXPENSES$"),
                CallbackQueryHandler(expenses_month_pick, pattern="^EXPENSE_MONTH$"),
                CallbackQueryHandler(expenses_month_show, pattern="^EXPENSE_MONTH_SHOW:"),
            ],
        },
        fallbacks=[CommandHandler("stop", stop)],
        allow_reentry=True,
    )

    app.add_handler(conv)

    # 🚀 Самый стабильный запуск webhook
    app.run_webhook(
        listen="0.0.0.0",
        port=port,
        url_path=WEBHOOK_PATH,
        webhook_url=webhook_url,
    )

    async def error_handler(update, context):
        print("🔥 ERROR:", context.error)
    
    app.add_error_handler(error_handler)

if __name__ == "__main__":
    main()


























































































































































































