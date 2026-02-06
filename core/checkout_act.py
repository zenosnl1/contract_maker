from docx import Document
from datetime import datetime
from docx.shared import Pt

from core.act_localization import (
    get_lang,
    yes_no,
    early_initiator,
    payment_method,
    safe,
)


# ======================================================
# MAIN
# ======================================================

def build_checkout_act(
    template_path: str,
    output_path: str,
    contract: dict,
    violations: list,
):

    doc = Document(template_path)

    lang = get_lang(contract)

    # ------------------------------
    # dates
    # ------------------------------

    start = datetime.fromisoformat(contract["start_date"]).date()

    end_plan = datetime.fromisoformat(contract["end_date"]).date()

    actual_raw = contract.get("actual_checkout_date")
    actual = (
        datetime.fromisoformat(actual_raw).date()
        if actual_raw
        else end_plan
    )

    # ------------------------------
    # nights
    # ------------------------------

    lived_nights = max(0, (actual - start).days)

    total_nights = contract.get("nights") or 0

    unused_nights = max(0, total_nights - lived_nights)

    # ------------------------------
    # finance
    # ------------------------------

    deposit = contract.get("deposit") or 0

    refund_unused_calc = unused_nights * (contract.get("price_per_day") or 0)

    refund_unused = contract.get("refund_unused_amount") or refund_unused_calc

    final_refund = contract.get("final_refund_amount") or 0

    extra_due = contract.get("extra_due_amount") or 0

    penalties_total = sum(v["amount"] for v in violations)

    # ------------------------------
    # localized values
    # ------------------------------

    values = {
        "CONTRACT_CODE": safe(contract.get("contract_code")),

        "FLAT_NUMBER": safe(contract.get("flat_number")),
        "CLIENT_NAME": safe(contract.get("client_name")),
        "CLIENT_ID": safe(contract.get("client_id")),
        "CLIENT_NUMBER": safe(contract.get("client_number")),
        "CLIENT_MAIL": safe(contract.get("client_mail")),
        "CLIENT_ADDRESS": safe(contract.get("client_address")),

        "START_DATE": start.strftime("%d.%m.%Y"),
        "END_DATE": end_plan.strftime("%d.%m.%Y"),
        "ACTUAL_CHECKOUT_DATE": actual.strftime("%d.%m.%Y"),
        "CHECKOUT_TIME": safe(contract.get("checkout_time")),

        "TOTAL_PRICE": safe(contract.get("total_price")),
        "DEPOSIT": safe(deposit),

        "EARLY_CHECKOUT": yes_no(contract.get("early_checkout"), lang),

        "EARLY_INITIATOR": early_initiator(
            contract.get("early_initiator"),
            lang,
        ),

        "EARLY_REASON": safe(contract.get("early_reason")),

        "REFUND_UNUSED_AMOUNT": safe(refund_unused),
        "FINAL_REFUND_AMOUNT": safe(final_refund),
        "EXTRA_DUE_AMOUNT": safe(extra_due),

        "LIVED_NIGHTS": str(lived_nights),
        "UNUSED_NIGHTS": str(unused_nights),

        "PENALTIES_TOTAL": safe(penalties_total),
    }

    replace_everywhere(doc, values)
    insert_violations_table(doc, violations, lang)

    doc.save(output_path)

    return output_path


# ======================================================
# HELPERS
# ======================================================

def replace_everywhere(doc: Document, data: dict):

    for p in doc.paragraphs:
        _process_paragraph(p, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _process_paragraph(p, data)


def _process_paragraph(p, data):

    text = p.text

    keys_used = [k for k in data if f"{{{{{k}}}}}" in text]

    if not keys_used:
        return

    for r in p.runs:
        r.text = ""

    i = 0
    while i < len(text):

        replaced = False

        for k in keys_used:

            ph = f"{{{{{k}}}}}"

            if text.startswith(ph, i):
                run = p.add_run(str(data[k]))
                run.font.size = Pt(11)
                run.bold = True
                i += len(ph)
                replaced = True
                break

        if not replaced:
            run = p.add_run(text[i])
            i += 1


# --------------------------------------------------


def insert_violations_table(doc: Document, violations: list, lang: str):

    marker = "{{VIOLATIONS_TABLE}}"

    for p in doc.paragraphs:

        if marker not in p.text:
            continue

        p.text = ""

        if not violations:

            txt = (
                "Нарушений не зафиксировано."
                if lang == "ru"
                else "Pārkāpumi nav konstatēti."
            )

            p.add_run(txt)
            return

        table = doc.add_table(rows=1, cols=3)

        hdr = table.rows[0].cells

        if lang == "ru":
            hdr[0].text = "Тип"
            hdr[1].text = "Сумма (€)"
            hdr[2].text = "Комментарий"
        else:
            hdr[0].text = "Veids"
            hdr[1].text = "Summa (€)"
            hdr[2].text = "Komentārs"

        for v in violations:
            row = table.add_row().cells
            row[0].text = v["violation_type"]
            row[1].text = str(v["amount"])
            row[2].text = v.get("description") or ""

        p._element.addnext(table._element)
        return
